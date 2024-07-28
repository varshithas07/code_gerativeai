from langchain.embeddings import HuggingFaceEmbeddings
from llama_index.embeddings.langchain import LangchainEmbedding
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import ChatPromptTemplate
from llama_index.core.llms import ChatMessage,MessageRole
from llama_index.llms.together import TogetherLLM
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader
import qdrant_client
from llama_index.vector_stores.qdrant import QdrantVectorStore
import re
from llama_index.core import StorageContext
from llama_index.core.postprocessor import SentenceTransformerRerank
# from llama_index.core.schema import MetadataMode
from llama_index.core.extractors import ( TitleExtractor,QuestionsAnsweredExtractor)
# from llama_index.extractors.entity import EntityExtractor
from llama_index.core.node_parser import TokenTextSplitter

from llama_index.core.ingestion import IngestionPipeline
from copy import deepcopy
from langchain_community.embeddings import HuggingFaceEmbeddings
from llama_index.core import Settings
from docx import Document
# from docx.shared import Pt
from docx import Document
# from llama_index.embeddings.langchain import LangchainEmbedding
from llama_index.core.postprocessor import SentenceTransformerRerank
# Create your views here.


rerank = SentenceTransformerRerank( model="cross-encoder/ms-marco-MiniLM-L-2-v2", top_n=7)

llm = TogetherLLM(
    model="togethercomputer/llama-2-70b-chat",
#     temperature= 0.7,
#      top_p= 0.8,
#      top_k=50,
api_key="f5e4f20a921125eaea37c452b325153c06adffbc76d649390b38b84445713485",
)



lc_embed_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-mpnet-base-v2"
)
embed_model = LangchainEmbedding(lc_embed_model)
Settings.embed_model = embed_model
# splitter = SentenceSplitter(chunk_size=1000)
Settings.llm =llm
# from llama_index.core import Settings
Settings.context_window = 3000

documents = SimpleDirectoryReader("./e_document").load_data()

text_splitter = TokenTextSplitter(
    separator=" ", chunk_size=1000, chunk_overlap=64
    )
title_extractor = TitleExtractor(nodes=40)
qa_extractor = QuestionsAnsweredExtractor(questions=3)
pipeline = IngestionPipeline(
 transformations=[text_splitter, title_extractor,qa_extractor])

nodes = pipeline.run(
    documents=documents,
    in_place=True,
    show_progress=True,)

#     # print(nodes[20].metadata)
# nodes_list = []
# for node in nodes:
#     # print(node)
#     nodes_list.append(node)

# # Optional: Print the list of nodes if needed
# print(nodes_list)
# title_questions_dict = {}

# for node in nodes:
#     title = node.metadata.get('document_title', 'Unknown Title').strip()
#     questions = node.metadata.get('questions_this_excerpt_can_answer', '').strip()
#     if questions:
#         if title not in title_questions_dict:
#             title_questions_dict[title] = []
#         title_questions_dict[title].append(questions)

# # Print the dictionary
# for title, questions in title_questions_dict.items():
#     print(f"Title: {title}")
#     for question in questions:
#         print(f" - Question: {question}")
#     print()

# Optionally, store the dictionary in a variable


# title_questions_dict = {}

# for node in nodes:
#     title = node.metadata.get('document_title', '').strip()
#     questions = node.metadata.get('questions_this_excerpt_can_answer', '').strip()

#     if questions:
#         # Extract questions from the formatted string
#         extracted_questions = [line.strip() for line in questions.split('\n') if line.strip().startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.')]
        
#         if title not in title_questions_dict:
#             title_questions_dict[title] = []
        
#         title_questions_dict[title].extend(extracted_questions)

# # Print the dictionary
# for title, questions in title_questions_dict.items():
#     print(f"Title: {title}")
#     for question in questions:
#         print(f" - Question: {question}")
#     print()
# print(title_questions_dict)

# # Example usage (assuming 'index' and 'text_qa_template' are already defined):
# title_questions_dict = {}

# for node in nodes:
#     title = node.metadata.get('document_title', '').strip()
#     questions = node.metadata.get('questions_this_excerpt_can_answer', '').strip()

#     if questions:
#         # Extract questions from the formatted string
#         extracted_questions = [line.strip() for line in questions.split('\n') if line.strip().startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.')]
        
#         if title not in title_questions_dict:
#             title_questions_dict[title] = []
        
#         title_questions_dict[title].extend(extracted_questions)

# title_questions_dict = {}

# for node in nodes:
#     title = node.metadata.get('document_title', '').strip()
#     summary = node.metadata.get('summary', '').strip()
#     questions = node.metadata.get('questions_this_excerpt_can_answer', '').strip()

#     # Extract questions from the formatted string
#     if questions:
#         extracted_questions = [line.strip() for line in questions.split('\n') if line.strip().startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.')]
        
#         # Add the extracted questions to the title_questions_dict
#         if title not in title_questions_dict:
#             title_questions_dict[title] = {'questions': [], 'summary': summary}
        
#         title_questions_dict[title]['questions'].extend(extracted_questions)

# # Process and save the extracted data
# for title, data in title_questions_dict.items():
#     # Process and format the questions
#     questions = '\n'.join(data['questions'])

title_questions_dict = {}

for node in nodes:
    title = node.metadata.get('document_title', '').strip()
    summary = node.metadata.get('summary', '').strip()
    questions = node.metadata.get('questions_this_excerpt_can_answer', '').strip()

    # Check if the title starts with "e-kranti" or comes after it alphabetically
    if title.lower() >= 'e-kranti':
        # Extract questions from the formatted string
        if questions:
            extracted_questions = [line.strip() for line in questions.split('\n') if line.strip().startswith('1.') or line.strip().startswith('2.') or line.strip().startswith('3.')]

            # Add the extracted questions to the title_questions_dict
            if title not in title_questions_dict:
                title_questions_dict[title] = {'questions': [], 'summary': summary}

            title_questions_dict[title]['questions'].extend(extracted_questions)

# Process and save the extracted data
for title, data in title_questions_dict.items():
    # Process and format the questions
    questions = '\n'.join(data['questions'])

    # Save the extracted data to the document or perform further processing as needed


# # Print the dictionary
# for title, questions in title_questions_dict.items():
#     print(f"Title: {title}")
#     for question in questions:
#         print(f" - Question: {question}")
#     print()
# # print(title_questions_dict)

client = qdrant_client.QdrantClient(path="qdrant1")

    # Create a local Qdrant vector store
vector_store = QdrantVectorStore(
client=client, collection_name="text_collection")
storage_context = StorageContext.from_defaults(vector_store=vector_store)
nodes_no_metadata = deepcopy(nodes)
# for node in nodes_no_metadata:
#      node.metadata = {
#         k: node.metadata[k]
#         for k in node.metadata
#         if k in ["page_label", "file_name"]
#     }
index= VectorStoreIndex(nodes=nodes_no_metadata,storage_context=storage_context)

chat_text_qa_msgs = [
                        ChatMessage(
                            role=MessageRole.SYSTEM,
                            content=(
                                 """You are an e-government online assistant chatbot system specifically developed for e-Governance Policy Initiatives under Digital India. 
            Your goal is to answer questions as accurately as possible based on the instructions and context provided."""
            " If the question is not related to the uploaded document, respond with 'I can only answer questions related to e-Governance Policy Initiatives under Digital India.'"
            " For general questions like 'Hi', 'How are you?', or 'Who are you?', respond accordingly, mentioning that you're here to assist with e-government policy using the context provided in the document and how you can help the user."
            " Given the context information and not prior knowledge."
            " Make sure to look for headings, subheadings, and key terms that match the question context."
                             
                            ),
                        ),
                        ChatMessage(
                            role=MessageRole.USER,
                            content=(
                                "Context information is below.\n"
                                "---------------------\n"
                                "{context_str}\n"
                                "---------------------\n"
                                "Given the context information and not prior knowledge, "
                                
                                "answer the question: {query_str} provided in bullet points or numbered list where appropriate.\n"
                                
                                

                               
                            ),
                        ),
                    ]
text_qa_template = ChatPromptTemplate(chat_text_qa_msgs)

# def ask_questions(index, text_qa_template):
#     while True:
#         # Ask the user for a question
#         question = input("Enter your question (or type 'no' to exit): ")

#         # Check if the user wants to exit
#         if question.lower() == 'no':
#             print("Exiting...")
#             break

#         # Query the index for the answer
#         response = index.as_query_engine(text_qa_template=text_qa_template).query(question)

#         # Print the response
#         print("Answer:", response)
#         print()

# # Call the function to start asking questions
# ask_questions(index, text_qa_template)



import os
# def clean_title(title_with_summary):
#     # Extract only the title from the given string
#     # Assuming title ends at the first period followed by a space
#     title = title_with_summary.split(". ")[0]
#     return title.strip()

# def ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict, output_file="output_new.txt"):
#     with open(output_file, 'w', encoding='utf-8') as file:
#         for title_with_summary, questions in title_questions_dict.items():
#             title = clean_title(title_with_summary)
#             file.write(f"Title: {title}\n")
#             file.write("=" * len(title) + "\n\n")  # Underline the title

#             for question in questions:
#                 # Query the index for the answer
#                 response = index.as_query_engine(text_qa_template=text_qa_template).query(question)
                
#                 # Write the question and response to the file
#                 file.write(f"Question: {question}\n")
#                 file.write(f"Answer: {response}\n\n")
#                 print(f"Processed question for title '{title}': {question}")

#             file.write("\n\n")  # Add space between sections

#     print(f"Results saved in file: {output_file}")

# # Example usage (assuming 'index' and 'text_qa_template' are already defined):
# ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict)

# def ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict, output_file="output_doc.txt"):
#     with open(output_file, 'w', encoding='utf-8') as file:
#         for title, questions in title_questions_dict.items():
#             file.write(f"Title: {title}\n")
#             file.write("=" * len(title) + "\n\n")  # Underline the title

#             for question in questions:
#                 # Query the index for the answer
#                 response = index.as_query_engine(text_qa_template=text_qa_template).query(question)
                
#                 # Write the question and response to the file
#                 file.write(f"Question: {question}\n")
#                 file.write(f"Answer: {response}\n\n")
#                 print(f"Processed question for title '{title}': {question}")

#             file.write("\n\n")  # Add space between sections

#     print(f"Results saved in file: {output_file}")

# def ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict, output_file="output_doc_1.txt"):
#     with open(output_file, 'w', encoding='utf-8') as file:
#         # Iterate through the title_questions_dict
#         for full_title, questions in title_questions_dict.items():
#             # Extract the title and summary
#             if '\n' in full_title:
#                 title, summary = full_title.split('\n', 1)
#             else:
#                 title = full_title
#                 summary = ""
                
#             # Write the title
#             file.write(f"# {title.strip()}\n\n")
            
#             # Write the summary
#             if summary.strip():
#                 file.write("## Summary\n\n")
#                 file.write(summary.strip() + "\n\n")
            
#             # Write the FAQ section
#             if questions:
#                 file.write("## FAQ\n\n")

#                 for question in questions:
#                     # Query the index for the answer
#                     response = index.as_query_engine(text_qa_template=text_qa_template).query(question)
                    
#                     # Write the question and response to the file
#                     file.write(f"**Question:** {question}\n\n")
#                     file.write(f"**Answer:** {response}\n\n")
#                     print(f"Processed question for title '{title}': {question}")

#                 file.write("\n")  # Add space between sections

#     print(f"Results saved in file: {output_file}")

# # Example usage (assuming 'index' and 'text_qa_template' are already defined):
# ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict)


# def ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict, output_file="output_doc_new_1.docx"):
#     # Create a new Document
#     doc = Document()

#     # Iterate through the title_questions_dict
#     for full_title, questions in title_questions_dict.items():
#         # Extract the title and summary
#         if '\n' in full_title:
#             title, summary = full_title.split('\n', 1)
#         else:
#             title = full_title
#             summary = ""
            
#         # Write the title
#         doc.add_heading(title.strip(), level=1)
        
#         # Write the summary
#         if summary.strip():
#             doc.add_heading('Summary', level=2)
#             doc.add_paragraph(summary.strip())
        
#         # Write the FAQ section
#         if questions:
#             doc.add_heading('FAQ', level=2)

#             for question in questions:
#                 # Query the index for the answer
#                 response = index.as_query_engine(text_qa_template=text_qa_template,similarity_top_k=3, node_postprocessors=[rerank]).query(question)
                
#                 # Extract the text from the response object
#                 response_text = str(response)  # Convert the response to a string
                
#                 # Write the question and response to the document
#                 doc.add_heading('Question:', level=3)
#                 doc.add_paragraph(question, style='List Bullet')
                
#                 doc.add_heading('Answer:', level=3)
#                 doc.add_paragraph(response_text)
                
#                 print(f"Processed question for title '{title}': {question}")

#             doc.add_paragraph()  # Add space between sections
#             doc.add_page_break()  # Add page break between titles

#     # Save the document
#     doc.save(output_file)
#     print(f"Results saved in file: {output_file}")

# ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict)


def ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict, output_file="output_doc_new_11.docx"):
    # Create a new Document
    doc = Document()

    # Iterate through the title_questions_dict
    for full_title, data in title_questions_dict.items():
        # Extract the title and summary
        if '\n' in full_title:
            title, summary = full_title.split('\n', 1)
        else:
            title = full_title
            summary = ""

        # Write the title
        doc.add_heading(title.strip(), level=1)

        # Write the summary
        if summary.strip():
            doc.add_heading('Summary', level=2)
            doc.add_paragraph(summary.strip())

        # Write the FAQ section
        if data['questions']:
            doc.add_heading('FAQ', level=2)

            for question in data['questions']:
                # Query the index for the answer
                response = index.as_query_engine(text_qa_template=text_qa_template,similarity_top_k=3, node_postprocessors=[rerank]).query(question)

                # Extract the text from the response object and remove initial content
                response_text = str(response).split('Answer:', 1)[-1].strip()

                # Write the question and response to the document
                doc.add_heading('Question:', level=3)
                doc.add_paragraph(question, style='List Bullet')

                doc.add_heading('Answer:', level=3)
                doc.add_paragraph(response_text)

                print(f"Processed question for title '{title}': {question}")

            doc.add_paragraph()  # Add space between sections
            doc.add_page_break()  # Add page break between titles

    # Save the document
    doc.save(output_file)
    print(f"Results saved in file: {output_file}")

ask_questions_and_save_to_single_file(index, text_qa_template, title_questions_dict)