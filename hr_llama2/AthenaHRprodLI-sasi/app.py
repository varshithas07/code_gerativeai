from flask import Flask, render_template, request, jsonify, session
from main import query, handle_greetings
app = Flask(__name__, static_folder='static',template_folder='templates')
app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'
from azure.storage.blob import BlobServiceClient
import os
import logging
import uuid
from docx import Document
import re
from flask_mail import Mail, Message
from flask import Flask
import io

#email configuration: to be changed from sasipriya004@gmail to the HR email/ whatever is the requirement.
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USERNAME'] = "sasipriya004@gmail.com"
app.config['MAIL_PASSWORD'] = "jbfq vund ytxl ipub"
app.config['MAIL_DEFAULT_SENDER'] = "sasipriya004@gmail.com"


mail = Mail(app)
#isolated per session, can also be adjusted accordingly if new chat feature is implemented. 
# per session means per user identifier
def main(input):
    if 'chat_history' not in session:
        session['chat_history'] = []
    if 'original_answers' not in session:
        session['original_answers'] = []  # New list to store original answers

    chat_history = session['chat_history']
    original_answers = session['original_answers']  # Retrieve the list of original answers

    if input == "exit":
        answer = "Goodbye!"
    elif handle_greetings(input):
        answer = "Hello! I am AthenaHR, here to assist you with any queries you may have related to Human resources within the company."
    else:
        response = query(input)
        answer = str(response)

        if answer is not None:
            formatted_answer = format_answer(answer)
            chat_history.append({"question": input, "answer": formatted_answer})
            original_answers.append({"question": input, "answer": answer})  # Store the original answer
            session['chat_history'] = chat_history  # Update the chat history in the session
            session['original_answers'] = original_answers  # Update the list of original answers
        else:
            answer = "No answer found"

    return answer, chat_history



#new filename is generated everytime user sends a email to save the chathistory in azure storage.
def generate_filename():
    return str(uuid.uuid4()) + ".docx"

import re
# def format_answer(answer):
#     lines = answer.split('\n')
#     formatted_lines = []
#     in_list = False
#     list_type = None
#     in_sublist = False

#     for line in lines:
#         # Check for numbered list
#         numbered_match = re.match(r'^(\d+\.\s)(.+)', line)
#         # Check for asterisk list
#         asterisk_match = re.match(r'^(\*\s)(.+)', line)
#         # Check for plus list
#         plus_match = re.match(r'^(\+\s)(.+)', line)
#         # Check for bold text
#         bold_match = re.match(r'\*\*(.*?)\*\*', line)

#         if numbered_match:
#             if in_list:  # Close the previous list
#                 formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')
#                 in_list = False
#             formatted_lines.append(f'<p>{line.strip()}</p>')
#         elif asterisk_match:
#             if not in_list:
#                 formatted_lines.append('<ul style="color: white;">')
#                 in_list = True
#                 list_type = 'ul'
#             formatted_lines.append(f'<li>{asterisk_match.group(2).strip()}</li>')
#         elif plus_match:
#             if not in_list:
#                 formatted_lines.append('<ul>')
#                 in_list = True
#                 list_type = 'ul'
#             if not in_sublist:
#                 formatted_lines.append('<ul>')
#                 in_sublist = True
#             formatted_lines.append(f'<li>{plus_match.group(2).strip()}</li>')
#         elif bold_match:
#             formatted_lines.append(f'<p><strong>{bold_match.group(1).strip()}</strong></p>')
#         else:
#             if in_list:  # Close the previous list
#                 if in_sublist:
#                     formatted_lines.append('</ul>')
#                     in_sublist = False
#                 formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')
#                 in_list = False
#             formatted_lines.append(f'<p>{line.strip()}</p>')

#     # Close any open lists
#     if in_sublist:
#         formatted_lines.append('</ul>')
#     if in_list:
#         formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')

#     # Combine all formatted lines
#     formatted_output = ''.join(formatted_lines)

#     return formatted_output

import re

def format_answer(answer):
    current_number = 1
    lines = answer.split('\n')
    formatted_lines = []
    in_list = False
    list_type = None

    for line in lines:
        # Check for numbered list
        numbered_match = re.match(r'^(\d+\.\s)(.+)', line)
        # Check for asterisk list
        asterisk_match = re.match(r'^(\*\s)(.+)', line)
        # Split asterisk list items that are on the same line
        asterisk_items = re.findall(r'\*\s(.+?)(?=(\*\s|$))', line)
        bold_match = re.match(r'^(\*\*)(.+?)(\*\*)', line)

        if bold_match:
                if not in_list or list_type != 'ol':
                    if in_list:  # Close the previous list
                        formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    formatted_lines.append('<ol>') 
                    in_list = True
                    list_type = 'ol'
                formatted_lines.append(f'<p style="font-weight: bold;">{current_number}. {bold_match.group(2).strip()}</p>')
                current_number += 1  # Increment current numbering
        
        elif numbered_match:
            if not in_list or list_type != 'ul':
                if in_list:  # Close the previous list
                    formatted_lines.append('</ol>' if list_type == 'ul' else '</ul>')
                formatted_lines.append('<ul>') 
                in_list = True
                list_type = 'ul'
            formatted_lines.append(f'<li style="margin-left: 50px;">{numbered_match.group(2).strip()}</li>')

        elif asterisk_match or asterisk_items:
            if not in_list or list_type != 'ul':
                if in_list:  # Close the previous list
                    formatted_lines.append('</ol>' if list_type == 'ol' else '</ul>')
                formatted_lines.append('<ul>')
                in_list = True
                list_type = 'ul'
            if asterisk_items:
                for item, _ in asterisk_items:
                    formatted_lines.append(f'<li style="margin-left: 50px;">{item.strip()}</li>')
            else:
                formatted_lines.append(f'{asterisk_match.group(2).strip()}</li>')
    

        else:
            if in_list:  # Close the previous list
                formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                in_list = False
            # Wrap non-list lines in paragraphs or handle them appropriately
            formatted_lines.append(f'<p>{line.strip()}</p>')

    # Close any open list tags
    if in_list:
        formatted_lines.append('</ul>' if list_type == 'ul' else '</ol>')

    # Combine all formatted lines
    formatted_output = ''.join(formatted_lines)

    return formatted_output

@app.route('/submit', methods=['POST'])
def handle_submit():
    input = request.form.get('user_input')
    answer, chat_history = main(input)
    print("chat:", chat_history)
    # session['chat_history'] = chat_history
    # session['original_answer'] = original_answer
    return render_template('athenachat.html', query=input, answer=answer, query_history=chat_history)


def save_chat_history_to_docx(email_hr_body):
        original_answers = session.get('original_answers', [])
        print(original_answers)
        # account_url = "https://csg100320027210ea93.blob.core.windows.net/"
        # connection_string = "W8KG2x7XjAVYqIxfGTSRmOF0adFJhD+2BwJ7dmF9IZhd1k6C5ZG7QcLK2Tsa6eP5cYu/iitHuXIK+AStxJMleQ=="
        connect_str = "DefaultEndpointsProtocol=https;AccountName=csg100320027210ea93;AccountKey=W8KG2x7XjAVYqIxfGTSRmOF0adFJhD+2BwJ7dmF9IZhd1k6C5ZG7QcLK2Tsa6eP5cYu/iitHuXIK+AStxJMleQ==;EndpointSuffix=core.windows.net"
        blob_service_client = BlobServiceClient.from_connection_string(connect_str)
        container_name = "hranalytics" 


        filename = generate_filename()

        doc = Document()

        for entry in original_answers:
            doc.add_paragraph(f"User: {entry['question']}")
            doc.add_paragraph(f"Chatbot: {entry['answer']}")
            doc.add_paragraph("")  # Add an empty paragraph for spacing

        doc.save('chat_history.docx')
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0) 
        
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=filename)
        blob_client.upload_blob(doc_io)

        subject = 'Athena Enquiry'
        recipient = ['sasipriya.a@quadwave.com', 'sumanyu.p@quadwave.com']

        # Build the email body
        email_body = f"Good evening HR \n\n {email_hr_body} \n\n Please find the chat history attached. \n\n Regards"

        with app.app_context():
            msg = Message(subject=subject, recipients=recipient, body=email_body)
            msg.attach("chat_history.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", doc_io.getvalue())
            mail.send(msg)

        # Delete the document from Azure Blob Storage
        blob_client.delete_blob()

        # Close the Document IO buffer
        doc_io.close()
    
@app.route('/save_history', methods=['POST'])
def save_history():
    # Ensure that the request method is POST
    if request.method == 'POST':
        # Parse JSON data from the request body
        data = request.json
        print("Request JSON Data:", data)
        if data:
            email_hr_body = data.get('email_hr_body')
            if email_hr_body:
                # Now you can use the email_hr_body data as needed
                save_chat_history_to_docx(email_hr_body)
                return jsonify(result="Email sent to HR successfully")
            else:
                return jsonify(error="Missing 'email_hr_body' field in JSON data"), 400
        else:
            return jsonify(error="No JSON data received"), 400
    else:
        return jsonify(error="Only POST requests are allowed"), 405

@app.route('/')
def index():
    return render_template('index.html',)
@app.route('/athena_chat')
def athena_chat():
    return render_template('athenachat.html')


if __name__ == '__main__':
    app.run(port = 7000, debug=False)
